import openpyxl
from django.http import HttpResponse
from .models import NonTeachingStaff,TeachingStaff,PhDStudent,Department, Course,ExamResult,ResearchProjectInformation,Year
from django.shortcuts import render
from openpyxl.styles import Alignment, Font, Border, Side
from django.shortcuts import render, redirect
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from .models import Course

from django.contrib.auth import authenticate, login
from django.shortcuts import render, redirect

def login_view(request):
    if request.method == "POST":
        username = request.POST["username"]
        password = request.POST["password"]
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            return redirect('dashboard')  # redirect to your dashboard
        else:
            # return error message if authentication fails
            return render(request, "login.html", {"login_error": "Wrong username or password"})
    return render(request, "login.html")


from django.contrib.auth.models import User
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from .models import (
    NonTeachingStaff, TeachingStaff, PhDStudent, Course, ExamResult,
    TeacherInformation, TeacherPublicationInformation, ResearchProjectInformation,
    SpecialNote, Specialpro, Special ,Year ,Department
)
from staff.form import (
    NonTeachingStaffForm, TeachingStaffForm, PhDStudentForm, CourseForm,
    ExamResultForm, TeacherInformationForm, TeacherPublicationInformationForm,
    ResearchProjectInformationForm, SpecialNoteForm, SpecialproForm, SpecialForm,YearForm,DepartmentForm,UserCreationCustomForm,UserEditCustomForm 
)

# Dashboard 
from django.contrib.auth.decorators import login_required

@login_required
def dashboard(request):
    user = request.user

    if user.is_superuser:
        # Show full dashboard
        context = {
            'years': Year.objects.all(),
            'departments': Department.objects.all(),
            'non_teaching_staff': NonTeachingStaff.objects.all(),
            'teaching_staff': TeachingStaff.objects.all(),
            'phd_students': PhDStudent.objects.all(),
            'courses': Course.objects.all(),
            'exam_results': ExamResult.objects.all(),
            'teacher_info': TeacherInformation.objects.all(),
            'teacher_publications': TeacherPublicationInformation.objects.all(),
            'research_projects': ResearchProjectInformation.objects.all(),
            'special_notes': SpecialNote.objects.all(),
            'specialpros': Specialpro.objects.all(),
            'specials': Special.objects.all(),
             'users': User.objects.all(), 
        }
    elif user.is_staff:
        # Show only the logged-in staff's related data
        context = {
            'teacher_info': TeacherInformation.objects.all(),
            'teacher_publications': TeacherPublicationInformation.objects.all(),
            'research_projects': ResearchProjectInformation.objects.all(),
            'special_notes': SpecialNote.objects.all(),
            'specialpros': Specialpro.objects.all(),
            'specials': Special.objects.all(),
        }
        return render(request, 'staff_dashboard.html', context)
    else:
        # Regular users see only courses and exam results, or you can show a message
        context = {
           
            # Or just an empty dict if you want to hide everything
        }

    return render(request, 'dashboard.html', context)


def add_user(request):
    if request.method == 'POST':
        form = UserCreationCustomForm(request.POST)
        if form.is_valid():
            user = form.save(commit=False)
            user.set_password(form.cleaned_data['password'])  # Hash password
            user.save()
            messages.success(request, "User created successfully.")
            return redirect('dashboard')
    else:
        form = UserCreationCustomForm()
    return render(request, 'edit_form.html', {'form': form, 'is_add': True})


@login_required
def edit_user(request, user_id):
    user_instance = get_object_or_404(User, id=user_id)
    if request.method == 'POST':
        form = UserEditCustomForm(request.POST, instance=user_instance)
        if form.is_valid():
            form.save()
            messages.success(request, "User updated successfully.")
            return redirect('dashboard')
    else:
        form = UserEditCustomForm(instance=user_instance)
    
    context = {
        'form': form,
        'is_add': False,
        'is_user_edit': True,
        'user_to_edit': user_instance,
    }
    return render(request, 'edit_form.html', context)


@login_required
def delete_user(request, user_id):
    user_instance = get_object_or_404(User, id=user_id)
    if request.method == 'POST':
        user_instance.delete()
        messages.success(request, "User deleted successfully.")
        return redirect('dashboard')
    return render(request, 'confirm_delete.html', {'object': user_instance, 'type': 'User'})
from django.contrib.auth.models import User
from django.contrib.auth.forms import AdminPasswordChangeForm
from django.contrib.auth import update_session_auth_hash
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required

@login_required
def change_user_password(request, user_id):
    user_instance = get_object_or_404(User, id=user_id)
    if request.method == 'POST':
        form = AdminPasswordChangeForm(user_instance, request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, f"Password updated successfully for {user_instance.username}.")
            # Important: If the user is the current logged in user, update session hash to keep them logged in
            if user_instance == request.user:
                update_session_auth_hash(request, form.user)
            return redirect('edit_user', user_id=user_instance.id)
    else:
        form = AdminPasswordChangeForm(user_instance)
    return render(request, 'change_password.html', {'form': form, 'user_obj': user_instance})


def add_year(request):
    return handle_crud(request, YearForm, 'dashboard')

def edit_year(request, pk):
    return handle_crud(request, YearForm, 'dashboard', instance=get_object_or_404(Year, pk=pk))

def delete_year(request, pk):
    return delete_instance(request, Year, pk)

def add_department(request):
    return handle_crud(request, DepartmentForm, 'dashboard')

def edit_department(request, pk):
    return handle_crud(request, DepartmentForm, 'dashboard', instance=get_object_or_404(Department, pk=pk))

def delete_department(request, pk):
    return delete_instance(request, Department, pk)

### NonTeachingStaff Views ###
def add_non_teaching_staff(request):
    return handle_crud(request, NonTeachingStaffForm, 'dashboard')

def edit_non_teaching_staff(request, pk):
    return handle_crud(request, NonTeachingStaffForm, 'dashboard', instance=get_object_or_404(NonTeachingStaff, pk=pk))

def delete_non_teaching_staff(request, pk):
    return delete_instance(request, NonTeachingStaff, pk)

### TeachingStaff Views ###
def add_teaching_staff(request):
    return handle_crud(request, TeachingStaffForm, 'dashboard')

def edit_teaching_staff(request, pk):
    return handle_crud(request, TeachingStaffForm, 'dashboard', instance=get_object_or_404(TeachingStaff, pk=pk))

def delete_teaching_staff(request, pk):
    return delete_instance(request, TeachingStaff, pk)

### PhDStudent Views ###
def add_phd_student(request):
    return handle_crud(request, PhDStudentForm, 'dashboard')

def edit_phd_student(request, pk):
    return handle_crud(request, PhDStudentForm, 'dashboard', instance=get_object_or_404(PhDStudent, pk=pk))

def delete_phd_student(request, pk):
    return delete_instance(request, PhDStudent, pk)

### Course Views ###
def add_course(request):
    return handle_crud(request, CourseForm, 'dashboard')

def edit_course(request, pk):
    return handle_crud(request, CourseForm, 'dashboard', instance=get_object_or_404(Course, pk=pk))

def delete_course(request, pk):
    return delete_instance(request, Course, pk)

### ExamResult Views ###
def add_exam_result(request):
    return handle_crud(request, ExamResultForm, 'dashboard')

def edit_exam_result(request, pk):
    return handle_crud(request, ExamResultForm, 'dashboard', instance=get_object_or_404(ExamResult, pk=pk))

def delete_exam_result(request, pk):
    return delete_instance(request, ExamResult, pk)

### TeacherInformation Views ###
def add_teacher_information(request):
    return handle_crud(request, TeacherInformationForm, 'dashboard')

def edit_teacher_information(request, pk):
    return handle_crud(request, TeacherInformationForm, 'dashboard', instance=get_object_or_404(TeacherInformation, pk=pk))

def delete_teacher_information(request, pk):
    return delete_instance(request, TeacherInformation, pk)

### TeacherPublicationInformation Views ###
def add_teacher_publication(request):
    return handle_crud(request, TeacherPublicationInformationForm, 'dashboard')

def edit_teacher_publication(request, pk):
    return handle_crud(request, TeacherPublicationInformationForm, 'dashboard', instance=get_object_or_404(TeacherPublicationInformation, pk=pk))

def delete_teacher_publication(request, pk):
    return delete_instance(request, TeacherPublicationInformation, pk)

### ResearchProjectInformation Views ###
def add_research_project(request):
    return handle_crud(request, ResearchProjectInformationForm, 'dashboard')

def edit_research_project(request, pk):
    return handle_crud(request, ResearchProjectInformationForm, 'dashboard', instance=get_object_or_404(ResearchProjectInformation, pk=pk))

def delete_research_project(request, pk):
    return delete_instance(request, ResearchProjectInformation, pk)

### SpecialNote Views ###
def add_special_note(request):
    return handle_crud(request, SpecialNoteForm, 'dashboard')

def edit_special_note(request, pk):
    return handle_crud(request, SpecialNoteForm, 'dashboard', instance=get_object_or_404(SpecialNote, pk=pk))

def delete_special_note(request, pk):
    return delete_instance(request, SpecialNote, pk)

### Specialpro Views ###
def add_specialpro(request):
    return handle_crud(request, SpecialproForm, 'dashboard')

def edit_specialpro(request, pk):
    return handle_crud(request, SpecialproForm, 'dashboard', instance=get_object_or_404(Specialpro, pk=pk))

def delete_specialpro(request, pk):
    return delete_instance(request, Specialpro, pk)

### Special Views ###
def add_special(request):
    return handle_crud(request, SpecialForm, 'dashboard')

def edit_special(request, pk):
    return handle_crud(request, SpecialForm, 'dashboard', instance=get_object_or_404(Special, pk=pk))

def delete_special(request, pk):
    return delete_instance(request, Special, pk)

### Common Utility Views ###
def handle_crud(request, form_class, redirect_url, instance=None):
    if request.method == 'POST':
        form = form_class(request.POST, instance=instance)
        if form.is_valid():
            form.save()
            messages.success(request, 'Saved successfully.')
            return redirect(redirect_url)
    else:
        form = form_class(instance=instance)
    return render(request, 'edit_form.html', {'form': form, 'is_add': instance is None})

def delete_instance(request, model_class, pk):
    instance = get_object_or_404(model_class, pk=pk)
    instance.delete()
    messages.success(request, 'Deleted successfully.')
    return redirect('dashboard')

def logout_view(request):
    logout(request)
    return redirect('login')

# Similar views for other models like TeachingStaff, PhDStudent, etc.
def year_selection(request):
    years = Year.objects.all()  # Fetch all years
    return render(request, 'year_selection.html', {'years': years})

def index(request):
    years = Year.objects.all().order_by('year')  # Get available years
    selected_year = request.GET.get('year', years.first().year if years else None)  # Default to the latest year if none selected

    context = {
        'years': years,
        'selected_year': selected_year,
    }
    return render(request, 'index.html', context)
  
import openpyxl
from django.http import HttpResponse
from .models import NonTeachingStaff, Year
from django.shortcuts import render
from openpyxl.styles import Alignment, Font, Border, Side

def year_selection(request):
    years = Year.objects.all()
    return render(request, 'year_selection.html', {'years': years})

def export_to_excel(request):
    selected_year = request.GET.get("year")  # Get selected year from request
    if not selected_year:
        return HttpResponse("Please select a year.", status=400)

    # Filter data based on selected year
    staff_data = NonTeachingStaff.objects.filter(year__year=selected_year)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = f"Non-Teaching Staff {selected_year}"
    title = f"Department wise Non-teaching Staff (Regular only) - {selected_year}"
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=9)
    title_cell = ws.cell(row=1, column=1, value=title)
    title_cell.font = Font(size=14, bold=True)
    title_cell.alignment = Alignment(horizontal="center", vertical="center")

    ws.append([""])
    ws.append(["Sr. No", "Name of Department", "Class 1", "", "Class 2", "", "Class 3", "", "Total"])
    ws.merge_cells(start_row=3, start_column=3, end_row=3, end_column=4)
    ws.merge_cells(start_row=3, start_column=5, end_row=3, end_column=6)
    ws.merge_cells(start_row=3, start_column=7, end_row=3, end_column=8)
    ws.append(["", "", "Female", "Male", "Female", "Male", "Female", "Male", ""])

    header_font = Font(bold=True)
    alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin")
    )

    for row in ws.iter_rows(min_row=3, max_row=4, min_col=1, max_col=9):
        for cell in row:
            cell.font = header_font
            cell.alignment = alignment
            cell.border = thin_border

    ws.column_dimensions['B'].width = 30  

    # Loop through staff data and add it to the Excel sheet
    for index, staff in enumerate(staff_data, start=1):
        row = [
            staff.sr_no,  # Serial Number
            staff.department.name if staff.department else "No Department",
            staff.class1_female, staff.class1_male,
            staff.class2_female, staff.class2_male,
            staff.class3_female, staff.class3_male,
            staff.total
        ]
        ws.append(row)

    for row in ws.iter_rows(min_row=5, max_row=ws.max_row, min_col=1, max_col=9):
        for cell in row:
            cell.alignment = alignment
            cell.border = thin_border

    note_row = ws.max_row + 2
    ws.cell(row=note_row, column=1, value="નોંધઃ જે વિભાગનું નામ ઉપર દર્શાવેલ ન હોય તેમ્ના વિભાગનું નામ ઉમેરવા વિનંતી")
    ws.merge_cells(start_row=note_row, start_column=1, end_row=note_row, end_column=9)
    note_cell = ws.cell(row=note_row, column=1)
    note_cell.alignment = Alignment(horizontal="left", vertical="center")
    note_cell.font = Font(italic=True)

    response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    response['Content-Disposition'] = f'attachment; filename="AR_Non-Teaching_Staff_{selected_year}.xlsx"'
    wb.save(response)
    return response

def teching(request):
    selected_year = request.GET.get("year")  # Get selected year from request
    if not selected_year:
        return HttpResponse("Please select a year.", status=400)

    # Debugging: Print selected year
    print(f"Selected Year: {selected_year}")

    # Filter data based on selected year
    staff_data = TeachingStaff.objects.filter(year__year=selected_year)

    # Debugging: Print count of records fetched
    print(f"Records fetched: {staff_data.count()}")

    if not staff_data.exists():
        return HttpResponse("No data available for the selected year.", status=404)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = f"Teaching Staff {selected_year}"

    title = f"Department-wise Teaching Staff - {selected_year}"
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=9)
    title_cell = ws.cell(row=1, column=1, value=title)
    title_cell.font = Font(size=14, bold=True)
    title_cell.alignment = Alignment(horizontal="center", vertical="center")

    ws.append([""])
    ws.append(["Sr. No", "Name of Department", "Professor", "", "Associate Professor", "", "Assistant Professor", "", "Total"])
    ws.merge_cells(start_row=3, start_column=3, end_row=3, end_column=4)
    ws.merge_cells(start_row=3, start_column=5, end_row=3, end_column=6)
    ws.merge_cells(start_row=3, start_column=7, end_row=3, end_column=8)
    ws.append(["", "", "Female", "Male", "Female", "Male", "Female", "Male", ""])

    header_font = Font(bold=True)
    alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin_border = Border(left=Side(style="thin"), right=Side(style="thin"), top=Side(style="thin"), bottom=Side(style="thin"))

    for row in ws.iter_rows(min_row=3, max_row=4, min_col=1, max_col=9):
        for cell in row:
            cell.font = header_font
            cell.alignment = alignment
            cell.border = thin_border

    ws.column_dimensions['B'].width = 30  # Adjust column width

    for index, staff in enumerate(staff_data, start=1):
        row = [
            staff.sr_no,
            staff.department.name if staff.department else "No Department",
            staff.professor_female, staff.professor_male,
            staff.associate_professor_female, staff.associate_professor_male,
            staff.assistant_professor_female, staff.assistant_professor_male,
            staff.total
        ]
        ws.append(row)

    for row in ws.iter_rows(min_row=5, max_row=ws.max_row, min_col=1, max_col=9):
        for cell in row:
            cell.alignment = alignment
            cell.border = thin_border

    note_row = ws.max_row + 2
    ws.cell(row=note_row, column=1, value="નોંધ: જો વિભાગનું નામ ઉપર દર્શાવેલ ન હોય, તો તેને ઉમેરવા વિનંતી.")
    ws.merge_cells(start_row=note_row, start_column=1, end_row=note_row, end_column=9)
    ws.cell(row=note_row, column=1).alignment = Alignment(horizontal="left", vertical="center")
    ws.cell(row=note_row, column=1).font = Font(italic=True)

    response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    response['Content-Disposition'] = f'attachment; filename="Teaching_Staff_{selected_year}.xlsx"'
    wb.save(response)
    return response


def PhD(request):
    selected_year = request.GET.get("year")
    if not selected_year:
        return HttpResponse("Please select a year.", status=400)

    # Filter PhD students based on selected year
    phd_students = PhDStudent.objects.filter(year__year=selected_year)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = f"PhD Students {selected_year}"

    title = f"Number of Ph.D. Students in University Departments - {selected_year}"
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=14)
    title_cell = ws.cell(row=1, column=1, value=title)
    title_cell.font = Font(size=14, bold=True)
    title_cell.alignment = Alignment(horizontal="center", vertical="center")

    ws.append([""])
    ws.append(["Sr. No", "Department Name", "Boys", "", "", "", "", "Girls", "", "", "", "", "Total", ""])
    ws.append(["", "", "General", "S.T", "S.C", "SEBC", "EWS", "General", "S.T", "S.C", "SEBC", "EWS", "Boys", "Girls"])

    ws.merge_cells(start_row=3, start_column=3, end_row=3, end_column=7)  # Boys
    ws.merge_cells(start_row=3, start_column=8, end_row=3, end_column=12)  # Girls
    ws.merge_cells(start_row=3, start_column=13, end_row=3, end_column=14)  # Total

    header_font = Font(bold=True)
    alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin")
    )

    for row in ws.iter_rows(min_row=3, max_row=4, min_col=1, max_col=14):
        for cell in row:
            cell.font = header_font
            cell.alignment = alignment
            cell.border = thin_border

    ws.column_dimensions['B'].width = 30

    for student in phd_students:
        row = [
            student.sr_no, 
            student.department.name if student.department else "No Department",

            student.boys_general, student.boys_st, student.boys_sc, student.boys_sebc, student.boys_ews,
            student.girls_general, student.girls_st, student.girls_sc, student.girls_sebc, student.girls_ews,
            student.total_boys, student.total_girls
        ]
        ws.append(row)

    for row in ws.iter_rows(min_row=5, max_row=ws.max_row, min_col=1, max_col=14):
        for cell in row:
            cell.alignment = alignment
            cell.border = thin_border

    note_row = ws.max_row + 2
    ws.cell(row=note_row, column=1, value="નોંધઃ જે વિભાગનું નામ ઉપર દર્શાવેલ ન હોય તેમ્ના વિભાગનું નામ ઉમેરવા વિનંતી.")
    ws.merge_cells(start_row=note_row, start_column=1, end_row=note_row, end_column=14)
    note_cell = ws.cell(row=note_row, column=1)
    note_cell.alignment = Alignment(horizontal="left", vertical="center")
    note_cell.font = Font(italic=True)

    response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    response['Content-Disposition'] = f'attachment; filename="AR_PhD_Students_{selected_year}.xlsx"'
    wb.save(response)
    return response
from django.http import HttpResponse
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side
from .models import Course, Department  # Ensure correct model import

def export(request):
    selected_year = request.GET.get("year")
    if not selected_year:
        return HttpResponse("Please select a year.", status=400)

    # Filter courses based on selected year
    students = Course.objects.filter(year__year=selected_year)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = f"Students Report {selected_year}"

    title = f"Number of Students in University Departments - {selected_year}"
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=15)
    title_cell = ws.cell(row=1, column=1, value=title)
    title_cell.font = Font(size=14, bold=True, color="0000FF")
    title_cell.alignment = Alignment(horizontal="center", vertical="center")

    # Merging header sections
    ws.merge_cells(start_row=2, start_column=4, end_row=2, end_column=8)
    ws.merge_cells(start_row=2, start_column=9, end_row=2, end_column=13)
    ws.merge_cells(start_row=2, start_column=14, end_row=2, end_column=15)

    ws.append(["Sr. No", "Department Name", "Course Name", "Boys", "", "", "", "", "Girls", "", "", "", "", "Total", ""])
    headers = [
        "General", "ST", "SC", "SEBC", "EWS",  # Boys
        "General", "ST", "SC", "SEBC", "EWS",  # Girls
        "Boys", "Girls"  # Total
    ]
    ws.append(["", "", ""] + headers)
    ws.column_dimensions['B'].width = 30  

    header_font = Font(bold=True)
    alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin_border = Border(left=Side(style="thin"), right=Side(style="thin"),
                         top=Side(style="thin"), bottom=Side(style="thin"))

    for row in ws.iter_rows(min_row=2, max_row=3, min_col=1, max_col=15):
        for cell in row:
            cell.font = header_font
            cell.alignment = alignment
            cell.border = thin_border

    row_num = 4
    sr_no = 1

    for department in Department.objects.all():
        department_courses = students.filter(department=department)
        first_row = True

        for course in department_courses:
            row = [
                sr_no if first_row else "",
                department.name if first_row else "",
                course.name,
                course.boys_general, course.boys_st, course.boys_sc, course.boys_sebc, course.boys_ews,
                course.girls_general, course.girls_st, course.girls_sc, course.girls_sebc, course.girls_ews,
                course.total_boys, course.total_girls
            ]
            ws.append(row)
            first_row = False
            row_num += 1

        if department_courses.exists():
            sr_no += 1  # Increment Sr. No only if department has results

    for row in ws.iter_rows(min_row=4, max_row=ws.max_row, min_col=1, max_col=15):
        for cell in row:
            cell.alignment = alignment
            cell.border = thin_border

    note_row = ws.max_row + 2
    ws.cell(row=note_row, column=1, value="Note: If a department name is missing, please ensure it is included.")
    ws.merge_cells(start_row=note_row, start_column=1, end_row=note_row, end_column=15)
    note_cell = ws.cell(row=note_row, column=1)
    note_cell.alignment = Alignment(horizontal="left", vertical="center")
    note_cell.font = Font(italic=True)

    response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    response['Content-Disposition'] = f'attachment; filename="AR_Students_{selected_year}.xlsx"'
    wb.save(response)

    return response

def export2(request):
    selected_year = request.GET.get("year")
    if not selected_year:
        return HttpResponse("Please select a year.", status=400)

    # Filter Exam Results based on selected year
    exam_results = ExamResult.objects.filter(year__year=selected_year)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = f"Exam Results {selected_year}"

    title = f"Exam Results by Department and Course - {selected_year}"
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=7)
    title_cell = ws.cell(row=1, column=1, value=title)
    title_cell.font = Font(size=14, bold=True)
    title_cell.alignment = Alignment(horizontal="center", vertical="center")

    ws.append([""])
    ws.append(["Sr. No", "Department Name", "Course Name", "First Class", "Second Class", "Pass Class", "Total Students"])

    # Header formatting
    header_font = Font(bold=True)
    alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin_border = Border(left=Side(style="thin"), right=Side(style="thin"),
                         top=Side(style="thin"), bottom=Side(style="thin"))

    for row in ws.iter_rows(min_row=3, max_row=3, min_col=1, max_col=7):
        for cell in row:
            cell.font = header_font
            cell.alignment = alignment
            cell.border = thin_border

    row_num = 4
    sr_no = 1

    for department in Department.objects.all():
        department_results = exam_results.filter(department=department)
        first_row = True

        for result in department_results:
            row = [
                sr_no if first_row else "",  # Sr. No
                department.name if first_row else "",  # Department Name
                result.course_name,
                result.first_class,
                result.second_class,
                result.pass_class,
                result.total_students,
            ]
            ws.append(row)
            first_row = False
            row_num += 1

        if department_results.exists():
            sr_no += 1  # Increment Sr. No only if department has results

    # Apply formatting for all rows
    for row in ws.iter_rows(min_row=4, max_row=ws.max_row, min_col=1, max_col=7):
        for cell in row:
            cell.alignment = alignment
            cell.border = thin_border

    # Add a footer note
    note_row = ws.max_row + 2
    ws.cell(row=note_row, column=1, value="નોંધઃ જે વિભાગનું નામ ઉપર દર્શાવેલ ન હોય તેમ્ના વિભાગનું નામ ઉમેરવા વિનંતી.")
    ws.merge_cells(start_row=note_row, start_column=1, end_row=note_row, end_column=7)
    note_cell = ws.cell(row=note_row, column=1)
    note_cell.alignment = Alignment(horizontal="left", vertical="center")
    note_cell.font = Font(italic=True)

    response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    response['Content-Disposition'] = f'attachment; filename="Exam_Results_{selected_year}.xlsx"'
    wb.save(response)

    return response


from django.http import HttpResponse
from docx import Document
from docx.shared import RGBColor, Pt
from .models import TeacherInformation, TeacherPublicationInformation, ResearchProjectInformation, SpecialNote, Specialpro, Special
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from io import BytesIO
from io import BytesIO
from django.http import FileResponse
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from .models import TeacherInformation 
from django.http import HttpResponse
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from io import BytesIO
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.enums import TA_CENTER
from datetime import datetime

def generate_pdf(request):
    selected_year = request.GET.get("year")
    if not selected_year:
        return HttpResponse("Please select a year.", status=400)

    buffer = BytesIO()
    pdf = SimpleDocTemplate(buffer, pagesize=A4)

    elements = []
    styles = getSampleStyleSheet()
    
    # **Title**
    title_style = styles['Title']
    title_style.alignment = TA_CENTER
    elements.append(Paragraph(f"Annual Report {selected_year} - Gujarat University", title_style))
    elements.append(Paragraph(f"અધ્યાપકોને લગતી માહિતી ({selected_year})", styles['Heading1']))

    # **Database Querying**
    teachers = TeacherInformation.objects.filter(year__year=selected_year)
    publications = TeacherPublicationInformation.objects.filter(year__year=selected_year)
    research_projects = ResearchProjectInformation.objects.filter(year__year=selected_year)
    special_note = SpecialNote.objects.filter(year__year=selected_year, is_displayed=True).first()
    special_pro = Specialpro.objects.filter(year__year=selected_year, is_displayed=True).first()
    special = Special.objects.filter(year__year=selected_year, is_displayed=True).first()

    # **Teacher Information Section**
    elements.append(Paragraph("(માત્ર વિદ્યાભવનો માટે) વિષયદીઠ રિફ્રેશર કોર્ષ/ઓરિયેન્ટેશન કોર્ષ/કાર્યશાળા/પરિસંવાદ વિષયક સંક્ષેપમાં માહિતી", styles['Normal']))

    # **Teacher Table**
    if teachers.exists():
        data = [['ક્રમ', 'શિક્ષકનું નામ અને હોદ્દો', 'કોન્ફરન્સ/પરિસંવાદ', 'સ્થાન અને તારીખ', 'રજુ કરેલા અભ્યાસ લેખનું શીર્ષક']]
        for idx, teacher in enumerate(teachers, start=1):
            data.append([str(idx), teacher.teacher_name, teacher.conference, teacher.location_date, teacher.paper_title])
        
        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold')
        ]))
        elements.append(table)
    else:
        elements.append(Paragraph("કોઈ માહિતી ઉપલબ્ધ નથી.", styles['Normal']))

    # **Publications Section**
    if publications.exists():
        elements.append(Paragraph("શિક્ષકના પ્રકાશનો (Publications):", styles['Heading2']))
        data = [['ક્રમ', 'શિક્ષકનું નામ', 'પ્રકાશન / અનુવાદિત પુસ્તક']]
        for index, publication in enumerate(publications, start=1):
            data.append([str(index), publication.teacher_name, publication.publication_title])

        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold')
        ]))
        elements.append(table)
    else:
        elements.append(Paragraph("કોઈ માહિતી ઉપલબ્ધ નથી.", styles['Normal']))

    # **Research Projects Section**
    if research_projects.exists():
        elements.append(Paragraph("શોધ પ્રોજેક્ટ (Research Projects):", styles['Heading2']))
        data = [['ક્રમ', 'શિક્ષકનું નામ', 'સંશોધન પ્રોજેક્ટનું શીર્ષક', 'સંશોધન ગ્રાંટ આપનાર સંસ્થા']]
        for index, project in enumerate(research_projects, start=1):
            data.append([str(index), project.teacher_name, project.research_project_title, project.funding_agency])

        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold')
        ]))
        elements.append(table)
    else:
        elements.append(Paragraph("કોઈ માહિતી ઉપલબ્ધ નથી.", styles['Normal']))

    # **Special Notes**
    elements.append(Paragraph(f"વિશેષ નોંધ: {special_note.note if special_note else 'માહિતી ઉપલબ્ધ નથી.'}", styles['Normal']))
    elements.append(Paragraph(f"વિશિષ્ટ પ્રયોગો-પ્રોજેક્ટો: {special_pro.note if special_pro else 'માહિતી ઉપલબ્ધ નથી.'}", styles['Normal']))
    elements.append(Paragraph(f"રમત-ગમત/યુવક મહોત્સવ/એન.એસ.એસ. નોંધપાત્ર વિશિષ્ટ સન્માન: {special.note if special else 'માહિતી ઉપલબ્ધ નથી.'}", styles['Normal']))

    # **Signature**
    elements.append(Paragraph("અધ્યક્ષશ્રીની સહિ અને સિક્કોઃ: ___________________________", styles['Normal']))

    pdf.build(elements)

    # Return PDF Response
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="Annual_Report_{selected_year}.pdf"'
    buffer.seek(0)
    response.write(buffer.read())
    buffer.close()
    return response
def download(request):
    selected_year = request.GET.get("year")  # Get selected year from request
    if not selected_year:
        return HttpResponse("Please select a year.", status=400)
    
    doc = Document()
    
    # **Header Section**
    header = doc.sections[0].header
    paragraph = header.paragraphs[0]
    paragraph.alignment = 0  # Left alignment
    run = paragraph.add_run(f'Gujarat University - Annual Report {selected_year}')
    run.font.size = Pt(12)

    # **Title Heading**
    heading = doc.add_heading('', level=1)
    run = heading.add_run(f'અધ્યાપકોને લગતી માહિતી ({selected_year})')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    heading.alignment = 1  # Center alignment

    # **Footer Section with Page Numbers**
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = 1  # Center alignment
    run = paragraph.add_run('વાર્ષિક અહેવાલ\n')
    run.font.size = Pt(10)
    run = paragraph.add_run("Page ")
    run.font.size = Pt(10)

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType', 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}space', 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType', 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType', 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

    # **Database Querying**
    teachers = TeacherInformation.objects.filter(year__year=selected_year)
    publications = TeacherPublicationInformation.objects.filter(year__year=selected_year)
    research_projects = ResearchProjectInformation.objects.filter(year__year=selected_year)
    special_note = SpecialNote.objects.filter(year__year=selected_year, is_displayed=True).first()
    special_pro = Specialpro.objects.filter(year__year=selected_year, is_displayed=True).first()
    special = Special.objects.filter(year__year=selected_year, is_displayed=True).first()

    teachers = TeacherInformation.objects.filter(year__year=selected_year)

    # **Teacher Information Section**
    doc.add_paragraph(style='List Bullet').add_run(
        '(માત્ર વિદ્યાભવનો માટે) વિષયદીઠ રિફ્રેશર કોર્ષ/ઓરિયેન્ટેશન કોર્ષ/કાર્યશાળા/પરિસંવાદ વિષયક સંક્ષેપમાં માહિતી'
    )

    # Create Table (Fixing column count to 5)
    teacher_table = doc.add_table(rows=1, cols=5)
    teacher_table.style = 'Table Grid'

    # Assign header text correctly
    header_cells = teacher_table.rows[0].cells
    header_cells[0].text = 'ક્રમ'  # Serial Number
    header_cells[1].text = 'શિક્ષકનું નામ અને હોદ્દો'  # Teacher Name & Designation
    header_cells[2].text = 'કોન્ફરન્સ/પરિસંવાદ'  # Conference/Seminar
    header_cells[3].text = 'સ્થાન અને તારીખ'  # Venue & Date
    header_cells[4].text = 'રજુ કરેલા અભ્યાસ લેખનું શીર્ષક'  # Research Paper Title

    # Populate Table with Data
    for idx, teacher in enumerate(teachers, start=1):
        row_cells = teacher_table.add_row().cells
        row_cells[0].text = str(idx)  # Serial Number
        row_cells[1].text = teacher.teacher_name
        row_cells[2].text = teacher.conference
        row_cells[3].text = teacher.location_date
        row_cells[4].text = teacher.paper_title 
        doc.add_paragraph("----" * 28)
    if publications.exists():
        doc.add_paragraph(style='List Bullet').add_run('શિક્ષકના પ્રકાશનો (Publications):')

        publication_table = doc.add_table(rows=1, cols=3)
        publication_table.style = 'Table Grid'
        header_cells = publication_table.rows[0].cells

        # Set table headers in Gujarati
        header_cells[0].text = 'ક્રમ'
        header_cells[1].text = 'શિક્ષકનું નામ'
        header_cells[2].text = 'પ્રકાશન / અનુવાદિત પુસ્તક'

        for index, publication in enumerate(publications, start=1):
            row_cells = publication_table.add_row().cells
            row_cells[0].text = str(index)  # Auto-generate the serial number
            row_cells[1].text = publication.teacher_name  # Fetch English field name
            row_cells[2].text = publication.publication_title  # Fetch English field name

    else:
        doc.add_paragraph("કોઈ માહિતી ઉપલબ્ધ નથી.", style='List Bullet') 

    doc.add_paragraph("----" * 28)

    # **Research Projects Section**
    if research_projects.exists():
        doc.add_paragraph(style='List Bullet').add_run('શોધ પ્રોજેક્ટ (Research Projects):')

        research_table = doc.add_table(rows=1, cols=4)
        research_table.style = 'Table Grid'
        header_cells = research_table.rows[0].cells

        # Gujarati Headers
        header_cells[0].text = 'ક્રમ'
        header_cells[1].text = 'શિક્ષકનું નામ'
        header_cells[2].text = 'સંશોધન પ્રોજેક્ટનું શીર્ષક'
        header_cells[3].text = 'સંશોધન ગ્રાંટ આપનાર સંસ્થા'

        for index, project in enumerate(research_projects, start=1):
            row_cells = research_table.add_row().cells
            row_cells[0].text = str(index)  # Auto-generate serial number
            row_cells[1].text = project.teacher_name  # Fetch data in English
            row_cells[2].text = project.research_project_title
            row_cells[3].text = project.funding_agency

    else:
        doc.add_paragraph("કોઈ માહિતી ઉપલબ્ધ નથી.", style='List Bullet')


    doc.add_paragraph("----" * 28)

    if special_note:
        doc.add_paragraph(style='List Bullet').add_run(f"વિશેષ નોંધ: \n {special_note.note}")
    else:
        doc.add_paragraph(style='List Bullet').add_run("વિશેષ નોંધ: માહિતી ઉપલબ્ધ નથી.")
    
    doc.add_paragraph("----" * 28)

    if special_pro:
        doc.add_paragraph(style='List Bullet').add_run(f"વિશિષ્ટ પ્રયોગો-પ્રોજેક્ટો: \n {special_pro.note}")
    else:
        doc.add_paragraph(style='List Bullet').add_run("વિશિષ્ટ પ્રયોગો-પ્રોજેક્ટો: માહિતી ઉપલબ્ધ નથી.")
    
    doc.add_paragraph("----" * 28)

    if special:
        doc.add_paragraph(style='List Bullet').add_run(f"રમત-ગમત/યુવક મહોત્સવ/એન.એસ.એસ. નોંધપાત્ર વિશિષ્ટ સન્માન: \n {special.note}")
    else:
        doc.add_paragraph(style='List Bullet').add_run("રમત-ગમત/યુવક મહોત્સવ/એન.એસ.એસ. નોંધપાત્ર વિશિષ્ટ સન્માન: માહિતી ઉપલબ્ધ નથી.")

    doc.add_paragraph("----" * 28)


    # **Signature Section**
    signature_paragraph = doc.add_paragraph()
    signature_paragraph.alignment = 2
    signature_paragraph.add_run('અધ્યક્ષશ્રીની સહિ અને સિક્કોઃ: ___________________________').font.size = Pt(12)

    # **Return Response as Word File**
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="Annual_Report_{selected_year}.docx"'
    doc.save(response)
    
    return response
  
